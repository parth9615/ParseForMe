import sys
import re
from pprint import pprint
from docx import Document
from docx.shared import Inches
import commands
import json
import datetime

def getRawData(filename):

    dictOfDatesAndInfo = {} # dictionary that maps from ('date' --> (eventType) , (time) , (description))
    rawListOfData = []
    weightDictionary = {}
    removeTableLineFromDocTable = False  # to remove | from doc tables

    if '.docx' in filename:                # if file is .docx then read from that method
        rawListOfData = readFromDOCX( filename)

    elif '.doc' in filename:                # if .doc file
        removeTableLineFromDocTable = True
        rawListOfData = readFromDOC(filename)

    elif '.txt' in filename:                # if .txt file
        rawListOfData = readFromTXT( filename)


    extractDates(dictOfDatesAndInfo, rawListOfData, removeTableLineFromDocTable , weightDictionary)
    return convertToJsonFormat(dictOfDatesAndInfo , weightDictionary)

#converts the dictionary to the  Json Dict formatting
def convertToJsonFormat (dictionary, weightDictionary):
    jsonList = []
    for key in dictionary.keys():
        weight = findWeightInDict(dictionary[key][0] , weightDictionary)

        if weight:
            SingleEventDict = {'Title' : dictionary[key][0] , 'Date' : key , 'Time' : dictionary[key][1] , 'Description' : dictionary[key][2], 'Weight' : weight  , 'Type': dictionary[key][0]}
        else:
            SingleEventDict = {'Title' : dictionary[key][0] , 'Date' : key , 'Time' : dictionary[key][1] , 'Description' : dictionary[key][2] , 'Weight' : int(0) , 'Type': dictionary[key][0]}

        jsonList.append(SingleEventDict)
    jsonDict = {"events" : jsonList}
    return jsonDict 


def findWeightInDict(eventType, weight):
    if eventType:
        for key in weight.keys():
            index = key.lower().find(eventType.lower())
            if index > -1:
                a = (re.findall(r'\d+.*?',key))

                return int(a[0])
    return None

def readFromDOCX(filename):
    rawListOfData= []
    document = Document(filename)       # open the document
    for paragraph in document.paragraphs:
        encodedText = paragraph.text.encode('ascii','ignore') # get encoded text from paragrpahs
        rawListOfData.append(encodedText)
    for table in document.tables:                             # get encoded text from tables
        for row in table.rows:
            encodedTableText = ''
            for cell in row.cells:
                encodedTableText += cell.text.encode('ascii','ignore') # get the text for that cell in proper formatting
            rawListOfData.append(encodedTableText)                     # add it to the list
    return rawListOfData

def readFromDOC(filename):

    newFileName = 'convertedFromDocToTxt.txt'                   # name of piping file
    cmd = 'antiword ' + filename + ' > ' + newFileName           # construct antiword command
    (status, output) = commands.getstatusoutput(cmd)
    if status:                                                  # if problem exit
        sys.stderr.write('there was an error: ' , output)
        sys.exit(1)
    else:
        return readFromTXT(newFileName)                                # else now readfromthetxt file

def readFromTXT( filename):
    f = open(filename, 'rU')              # Open and read the file. for read only
    rawListOfData = f.readlines()         # get each line as a list
    return rawListOfData

def extractDates(dictOfDatesAndInfo, rawListOfData, removeTableLineFromDocTable, weightDictionary):
  dayAndMonthList = ['Monday' , 'Tuesday' , 'Wednesday' , 'Thursday' , 'Friday', 'Saturday' , 'Sunday'
'Mondays' , 'Tuesdays' , 'Wednesdays' , 'Thursdays' , 'Fridays' ,'Saturdays' , 'Sundays' 'Mon' , 'Tue',
 'Wed' , 'Thur' , 'Fri' , 'Sat' , 'Sun' , 'January', 'February' , 'March' , 'April' , 'May' , 'June',
'July' , 'August' , 'September' , 'October' , 'November' , 'December' ,'Jan' , 'Feb' , 'Mar',
   'Apr' , 'May' , 'Jun' , 'Jul' , 'Aug' , 'Sep' , 'Oct' , 'Nov' , 'Dec'  ]
  relevantDates = []                    # make a list to hold all the dates
                                       # get all days and store in relevantDates


  extractDays(dictOfDatesAndInfo,relevantDates , rawListOfData, dayAndMonthList, removeTableLineFromDocTable, weightDictionary)

'''
This method iterates through the rawListOfData to find patterns in the method
Current parsing methods included in the pattern
1: Finding a line that contains a day of the week
2: Finding a line that contains a time in the format ##:## (for ex 12:10 or 3:20)
3: Finding a line that contains a month of the year
4: Finding a line that contains a date in the format ##/##
todo : update this numeric list as more patterns are added:
'''
def extractDays(dictOfDatesAndInfo, relevantDates , rawListOfData, dayAndMonthList, removeTableLineFromDocTable, weightDictionary):
    optimizationBoolean = False
    for individualLine in rawListOfData:  # iterate through each line
        ##print(individualLine)
        individualLine = individualLine.lower()
        if removeTableLineFromDocTable:
            individualLine = individualLine.replace("|" , "") # remove table lines
        for days in dayAndMonthList:              # iterate through each day combination
            # # regex pattern to find the entire line that contains a day in the dayList
            # dayOfTheWeekPattern = ('.+')+(days)+('?\s.+')
            # # regex flag to ingnorecase and make the dot include whitespace
            # dayOfTheWeekFlags =   re.IGNORECASE | re.DOTALL
            # # call findInString to check if such a pattern exists
            # result = findInString(dayOfTheWeekPattern , individualLine , dayOfTheWeekFlags, relevantDates)
            # if result:
             if days.lower() in individualLine:
                # print(individualLine , "HELP")
                makeEventFromMonth(individualLine, dictOfDatesAndInfo)
                break

        extractWeight(individualLine , weightDictionary)
        #pattern to find the pattern ##/## which is commonly used to denote dates
        dateTimePattern = ('\d\d?[\./-]\d\d?')
        dateTimeFlags   = re.DOTALL
        result = findInString(dateTimePattern , individualLine , dateTimeFlags , relevantDates)
        if result:
            getDateSeperatingIndex = getDateSeperatingChar(result , ['.' , '/' , '-'])

            if getDateSeperatingIndex:
                month = result[0:getDateSeperatingIndex]
                day = result[getDateSeperatingIndex+1:]

                result =  month + '/' +day
                result += getValidYear(result)
            #checkDateValidity(result)
            if (checkIfDateIsValid(result)):
                makeEventFromDate(result , individualLine , dictOfDatesAndInfo)

# def checkDateValidity(date_text):
#     try:
#         datetime.datetime.strptime(date_text, '%m/%d')
#         print 'VALID DATE ))))))))))))' , date_text
#     except ValueError:
#         print ValueError("Incorrect data format, should be YYYY-MM-DD was " , date_text)


def getDateSeperatingChar(searchString, Charlist):
    for char in Charlist:
        index = searchString.find(char)
        if index > -1:
            return index

def extractWeight(lineToSearch, weightDictionary):
    findPercentIndex = lineToSearch.find('%')  # if percent sign is present
    if findPercentIndex > -1:                  # valid index found
        findNumericalWeight = re.search('\d?\d?\.?\d?\d' , lineToSearch)  # find any possible numerical percent
        if findNumericalWeight:
            weightDictionary[lineToSearch] = findNumericalWeight.group()


def makeEventFromDate(date, stringToSearch, dictionary):

    #get the proper event and time and then add to dictionary
    eventType = getEventType(stringToSearch)
    if eventType:

        time =      getValidTime(stringToSearch) # time returned in (firstTimeFound) , (secondTimeFound) ,(firstTimeFound , 'to' , secondTimeFound)
        dictionary[date] = (eventType) , (time) , (stringToSearch)


def getInfo(event,  time, date, stringToSearch):

    removeStuffList = [event, time, date]
    # removedEvent = removeSubstring(stringToSearch , event) # return string with removed event
    # if time:                                               # if valid time was there
    #     rmTime0 = removeSubstring(removedEvent , time) #rm both times
    #     removedDate = removeSubstring(rmTime0, date)
    # else :
    #     removedDate = removeSubstring(removedEvent, date) #  return string with removed event, time, date
    #     return removedDate

    return " ".join( (removeSubstring(stringToSearch , removeStuffList)).split())



def removeSubstring(stringToSearch, listToFind):
    for individualItems in listToFind:
        if isinstance(individualItems , str):                 # if findingObject is of type string
            stringToSearch = stringToSearch.replace(individualItems , '') # return the removed version

    return stringToSearch                             # else return the original string

def makeEventFromMonth(stringToSearch, dictOfDatesAndInfo):

    monthFound = ''
    dateFound = ''
    informationArroundDate = ''
    monthToNumDict = {'January' :1, 'February':2 , 'March':3 , 'April':4 , 'May':5, 'June' :6,
   'July':7 , 'August':8 , 'September':9 , 'October':10 , 'November':11 , 'December':12 ,'Jan':1 ,
    'Feb':2 , 'Mar':3, 'Apr':4 , 'May':5 , 'Jun':6 , 'Jul':7 , 'Aug':8 , 'Sep':9 , 'Oct':10 ,
     'Nov':11 , 'Dec':12}
    for month in monthToNumDict.keys():
        # monthFound = re.search ('\s' + month , stringToSearch, re.IGNORECASE) # try to find a month
        monthFound = month.lower() in  stringToSearch
        if monthFound:

            monthIndex = stringToSearch.lower().find(month.lower()) # find the index of where the month is
            # print(stringToSearch , month)
            # print('stringLenght =' , len(stringToSearch) , 'monthIndex' , monthIndex , 'monthIndex + 10' , monthIndex+10 , 'monthIndex-10' , monthIndex-10)
            if ((monthIndex+30 >= len(stringToSearch)) or (monthIndex-30 <= 0)):
                dateFound = re.search( '\d\d?' , stringToSearch)

            else:
                # try to find a date with a range of +- 10 chars from where the monthIndex was
                dateFound = re.search( '\d\d?' , stringToSearch[monthIndex - 10: monthIndex+10])


            #print 'the month = =====' , month , stringToSearch
            if dateFound:

                # make date in proper format of month/day
                dateFoundFormatted = str(monthToNumDict[month]) + '/' + dateFound.group()
                dateFoundFormatted += getValidYear(dateFoundFormatted)
                if checkIfDateIsValid(dateFoundFormatted):
                    # now add to the dictionary with date found
                    makeEventFromDate(dateFoundFormatted , stringToSearch , dictOfDatesAndInfo)
                break


def getValidYear(stringWithoutYear):
    dateNow = datetime.datetime.now()
    year = ''
    if ((dateNow.month + 6) % 12) > 6:
        year += '/' + str(dateNow.year + 1)
    else:
        year += '/' + str(+ dateNow.year)
    return year

def checkIfDateIsValid(dateString):
    dateSplit = re.findall(r'\d+',dateString)

    try:
        datetime.date(year = int(dateSplit[2]), month = int(dateSplit[0]), day = int(dateSplit[1]) )
    except ValueError:
        return False
    return True

def getEventType(stringToSearch):
    events = [ 'midterm' , 'final exam' ,'final' , 'test' , 'quiz' ,  'assignment'
    , 'assignments' , 'exam' , 'homework'  , 'paper' , 'essay' , 'project' ]
    for event in events:
        eventIndex = stringToSearch.lower().find(event) # find index of event
        if eventIndex > -1:                             # if event exists
            return event


# return a tuple in form (firstTimeFound) , (secondTimeFound) ,(firstTimeFound , 'to' , secondTimeFound)
def getValidTime(stringToSearch):
    firstTimeFound = getTime(stringToSearch)
    if firstTimeFound:
        firstTimeIndex = stringToSearch.find(firstTimeFound)
        # try to get a second time to establish valid time frame:
        secondTimeFound = getTime(stringToSearch[firstTimeIndex + len(firstTimeFound):])
        if secondTimeFound:
            return (firstTimeFound , 'to' , secondTimeFound)  # return time in this format
        else:
            return  (firstTimeFound)
    else:
        # time in above format not found try searching for am/pm
        am = stringToSearch.lower().find('am')
        if am > -1:
            time = re.search('\d\d?:?\d?\d?' , stringToSearch)
            if time:
                return time.group()
        pm = stringToSearch.lower().find('pm')
        if pm > -1:
            time = re.search('\d\d?.\d?\d?' , stringToSearch)
            if time:
                return time.group()

def getTime(stringToSearch):
    # pattern to find the pattern ##:## which is commonly used to denote time
    clockTimePattern = ('\d\d?:\d\d?')
    # flag to make the dot include whitespace
    clockTimeFlags   = re.DOTALL
    # return a time if found:
    timeFound = re.search(clockTimePattern , stringToSearch , clockTimeFlags)
    if timeFound:
        return timeFound.group()

'''
patternToFind = regexExpression that the regular expression looks for
textToSearch  = the text to find the pattern in
flags         = any and all flags required for the paticular search
relevantDates = the list to store the line if pattern is matched
This method takes the patternToFind and tries to search for it in the
textToSearch and if it is found then it stores it in the relevantDates if
it already doesn't exist
'''
def findInString(patternToFind, textToSearch, flags, relevantDates):
    if flags:                                      # if flags not empty
        matchResult = re.search((patternToFind) , textToSearch, flags)
    else:                                          # no flags to include
        matchResult = re.search((patternToFind) , textToSearch)

    if matchResult:                                  # if there was a match
        if matchResult.group() not in relevantDates: # if match doesn't exist in datex
            relevantDates.append(matchResult.group()) # add it to the date list
            return matchResult.group()

def main():

  # This command-line parsing code is provided.
  # Make a list of command line arguments, omitting the [0] element
  # which is the script itself.
  args = sys.argv[1:]

  for filename in args:                 # iterate through each filename
    getRawData(filename)     # extract dates from each file



## boiler plate python main function
if __name__ == '__main__':
  main()
