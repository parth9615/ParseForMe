import sys
import re
from pprint import pprint
from docx import Document
from docx.shared import Inches
import commands

def getRawData(filename):


    dictOfDatesAndInfo = {} # dictionary that maps from ('date' --> (eventType) , (time) , (description))
    rawListOfData = []

    if '.docx' in filename:                # if file is .docx then read from that method
        rawListOfData = readFromDOCX( filename)

    elif '.doc' in filename:                # if .doc file
        rawListOfData = readFromDOC(filename)

    elif '.txt' in filename:                # if .txt file
        rawListOfData = readFromTXT( filename)
    
    extractDates(dictOfDatesAndInfo, rawListOfData)
    pprint(dictOfDatesAndInfo)                     # print the result


def readFromDOCX(filename):
    rawListOfData= []
    document = Document(filename)       # open the document
    for paragraph in document.paragraphs:
        encodedText = paragraph.text.encode('ascii','ignore') # get encoded text from paragrpahs
        rawListOfData.append(encodedText)
    for table in document.tables:                             # get encoded text from tables
        for row in table.rows:
            encodedTableText = ''
            for cell in row.cells:
                encodedTableText += cell.text.encode('ascii','ignore') # get the text for that cell in proper formatting
            rawListOfData.append(encodedTableText)                     # add it to the list
    return rawListOfData

def readFromDOC(filename):

    newFileName = 'convertedFromDocToTxt.txt'                   # name of piping file
    cmd = 'antiword ' + filename + ' > ' + newFileName           # construct antiword command
    (status, output) = commands.getstatusoutput(cmd)
    if status:                                                  # if problem exit
        sys.stderr.write('there was an error: ' , output)
        sys.exit(1)
    else:
        return readFromTXT(newFileName)                                # else now readfromthetxt file

def readFromTXT( filename):
    f = open(filename, 'rU')              # Open and read the file. for read only
    rawListOfData = f.readlines()         # get each line as a list
    return rawListOfData

def extractDates(dictOfDatesAndInfo, rawListOfData):
  dayAndMonthList = ['Monday' , 'Tuesday' , 'Wednesday' , 'Thursday' , 'Friday', 'Saturday' , 'Sunday'
'Mondays' , 'Tuesdays' , 'Wednesdays' , 'Thursdays' , 'Fridays' ,'Saturdays' , 'Sundays' 'Mon' , 'Tue',
 'Wed' , 'Thur' , 'Fri' , 'Sat' , 'Sun' , 'January', 'February' , 'March' , 'April' , 'May' , 'June',
'July' , 'August' , 'September' , 'October' , 'November' , 'December' ,'Jan' , 'Feb' , 'Mar',
   'Apr' , 'May' , 'Jun' , 'Jul' , 'Aug' , 'Sep' , 'Oct' , 'Nov' , 'Dec'  ]
  relevantDates = []                    # make a list to hold all the dates
                                       # get all days and store in relevantDates


  extractDays(dictOfDatesAndInfo,relevantDates , rawListOfData, dayAndMonthList)

'''
This method iterates through the rawListOfData to find patterns in the method
Current parsing methods included in the pattern
1: Finding a line that contains a day of the week
2: Finding a line that contains a time in the format ##:## (for ex 12:10 or 3:20)
3: Finding a line that contains a month of the year
4: Finding a line that contains a date in the format ##/##
todo : update this numeric list as more patterns are added:
'''
def extractDays(dictOfDatesAndInfo, relevantDates , rawListOfData, dayAndMonthList):
    optimizationBoolean = False
    for individualLine in rawListOfData:  # iterate through each line
        for days in dayAndMonthList:              # iterate through each day combination
            # regex pattern to find the entire line that contains a day in the dayList
            dayOfTheWeekPattern = ('.+')+(days)+('\s.+')
            # regex flag to ingnorecase and make the dot include whitespace
            dayOfTheWeekFlags =   re.IGNORECASE | re.DOTALL
            # call findInString to check if such a pattern exists
            result = findInString(dayOfTheWeekPattern , individualLine , dayOfTheWeekFlags, relevantDates)
            if result:
                makeEventFromMonth(individualLine, dictOfDatesAndInfo)

                break

        # pattern to find the pattern ##/## which is commonly used to denote dates
        dateTimePattern = ('\d\d?/\d\d?')
        dateTimeFlags   = re.DOTALL
        result = findInString(dateTimePattern , individualLine , dateTimeFlags , relevantDates)
        if result:
             makeEventFromDate(result , individualLine , dictOfDatesAndInfo)



def makeEventFromDate(date, stringToSearch, dictionary):

    #get the proper event and time and then add to dictionary
    eventType = getEventType(stringToSearch)
    time =      getValidTime(stringToSearch)
    if time:                                    # if time found send the formatted version
        dictionary[date] = (eventType) , (time[2]) , (getInfo(eventType, time , date, stringToSearch))
    else:
        dictionary[date] = (eventType) , (time) , (getInfo(eventType, time , date, stringToSearch))

def getInfo(event,  time, date, stringToSearch):
    removedEvent = removeSubstring(stringToSearch , event) # return string with removed event
    if time:                                               # if valid time was there
        rmTime0 = removeSubstring(removedEvent , time[0]) #rm both times
        rmTime1 = removeSubstring(rmTime0 , time[1])
        removedDate = removeSubstring(rmTime1, date)
    else :
        removedDate = removeSubstring(removedEvent, date) #  return string with removed event, time, date
    return removedDate                                # return string with printed info deleted


def removeSubstring(stringToSearch, findingObject):
    if isinstance(findingObject , str):                 # if findingObject is of type string
        return stringToSearch.replace(findingObject , '') # return the removed version
    else:
        return stringToSearch                             # else return the original string

def makeEventFromMonth(stringToSearch, dictOfDatesAndInfo):
    monthFound = ''
    dateFound = ''
    informationArroundDate = ''
    monthToNumDict = {'January' :1, 'February':2 , 'March':3 , 'April':4 , 'May':5, 'June' :6,
   'July':7 , 'August':8 , 'September':9 , 'October':10 , 'November':11 , 'December':12 ,'Jan':1 ,
    'Feb':2 , 'Mar':3, 'Apr':4 , 'May':5 , 'Jun':6 , 'Jul':7 , 'Aug':8 , 'Sep':9 , 'Oct':10 ,
     'Nov':11 , 'Dec':12}
    for month in monthToNumDict.keys():
        monthFound = re.search ('\s' + month , stringToSearch, re.IGNORECASE) # try to find a month
        if monthFound:

            monthIndex = stringToSearch.lower().find(month.lower()) # find the index of where the month is
            # try to find a date with a range of +- 10 chars from where the monthIndex was
            dateFound = re.search( '\d\d?' , stringToSearch[monthIndex - 10: monthIndex+10])
            #print 'the month = =====' , month , stringToSearch
            if dateFound:
                # make date in proper format of month/day
                dateFoundFormatted = str(monthToNumDict[month]) + '/' + dateFound.group()

                # now add to the dictionary with date found
                makeEventFromDate(dateFoundFormatted , stringToSearch , dictOfDatesAndInfo)
                break



def getEventType(stringToSearch):
    events = ['final exam' , 'test' , 'quiz' , 'office-hours' , 'office hours' , 'assignment'
    , 'assignments' , 'exam' , 'homework' , 'webassign']
    for event in events:
        eventIndex = stringToSearch.lower().find(event) # find index of event
        if eventIndex > -1:                             # if event exists
            return event



def getValidTime(stringToSearch):
    firstTimeFound = getTime(stringToSearch)
    if firstTimeFound:
        firstTimeIndex = stringToSearch.find(firstTimeFound)
        # try to get a second time to establish valid time frame:
        secondTimeFound = getTime(stringToSearch[firstTimeIndex + len(firstTimeFound):])
        if secondTimeFound:
            return (firstTimeFound) , (secondTimeFound) ,(firstTimeFound , 'to' , secondTimeFound)  # return time in this format


def getTime(stringToSearch):
    # pattern to find the pattern ##:## which is commonly used to denote time
    clockTimePattern = ('\d\d?:\d\d?')
    # flag to make the dot include whitespace
    clockTimeFlags   = re.DOTALL
    # return a time if found:
    timeFound = re.search(clockTimePattern , stringToSearch , clockTimeFlags)
    if timeFound:
        return timeFound.group()

'''
patternToFind = regexExpression that the regular expression looks for
textToSearch  = the text to find the pattern in
flags         = any and all flags required for the paticular search
relevantDates = the list to store the line if pattern is matched
This method takes the patternToFind and tries to search for it in the
textToSearch and if it is found then it stores it in the relevantDates if
it already doesn't exist
'''
def findInString(patternToFind, textToSearch, flags, relevantDates):
    if flags:                                      # if flags not empty
        matchResult = re.search((patternToFind) , textToSearch, flags)
    else:                                          # no flags to include
        matchResult = re.search((patternToFind) , textToSearch)

    if matchResult:                                  # if there was a match
        if matchResult.group() not in relevantDates: # if match doesn't exist in datex
            relevantDates.append(matchResult.group()) # add it to the date list
            return matchResult.group()

def main():

  # This command-line parsing code is provided.
  # Make a list of command line arguments, omitting the [0] element
  # which is the script itself.
  args = sys.argv[1:]

  for filename in args:                 # iterate through each filename
    getRawData(filename)     # extract dates from each file



## boiler plate python main function
if __name__ == '__main__':
  main()
